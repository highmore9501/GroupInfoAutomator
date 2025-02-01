from docx import Document
from typing import Dict, Any, List, Optional
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from datetime import datetime


def get_trip_slice(attraction_info: List[Dict[str, Any]], start_date: str, trip_days: int) -> List[Dict[str, Any]]:
    """
    根据给定的日期和行程天数，返回景点信息的全部或者切片
    """
    # 找到给定日期的索引
    start_index = next((i for i, item in enumerate(
        attraction_info) if item['日期'] == start_date), None)

    if start_index is None:
        raise ValueError("给定的日期不在景点信息中")

    # 如果行程天数大于景点信息里的元素数量，返回整个景点信息
    if trip_days >= len(attraction_info):
        return attraction_info

    # 如果行程天数小于景点信息里的元素数量，返回包含给定日期的连续切片，长度为行程天数
    end_index = start_index + trip_days
    if end_index >= len(attraction_info):
        start_index = len(attraction_info) - trip_days
        return attraction_info[start_index:]
    else:
        return attraction_info[start_index:end_index]


def count_guests(guests):
    """
    计算客人数量
    """
    adult_count = 0
    child_count = 0
    current_year = datetime.now().year

    for guest in guests:
        guest_id = guest['id']
        guest_year = int(guest_id[6:10])
        if current_year - guest_year >= 18:
            adult_count += 1
        else:
            child_count += 1

    return adult_count, child_count


def replace_in_xlsx(wb_path, replacements):
    """
    遍历整个工作簿并替换指定的内容
    """
    wb = load_workbook(wb_path)

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for old_value, new_value in replacements.items():
                        if old_value in cell.value:
                            cell.value = cell.value.replace(
                                old_value, new_value)

    wb.save(wb_path)


def copy_xlsx_row_and_insert(wb_path, sheet_name, source_row_index, copy_count):
    """
    复制源行并插入到表格中
    """
    wb = load_workbook(wb_path)
    sheet = wb[sheet_name]
    source_row = list(sheet.iter_rows(min_row=source_row_index,
                      max_row=source_row_index, values_only=False))[0]

    for i in range(copy_count):
        new_row_index = source_row_index + i + 1
        sheet.insert_rows(new_row_index)

        for col_index, cell in enumerate(source_row, start=1):
            new_cell = sheet.cell(row=new_row_index, column=col_index)
            new_cell.value = cell.value
            if cell.has_style:
                new_cell._style = cell._style
            if cell.hyperlink:
                new_cell.hyperlink = cell.hyperlink
            if cell.comment:
                new_cell.comment = cell.comment

    wb.save(wb_path)


def searchSight(sight: str, attraction_info: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    在行程的景点信息中搜索指定的景点信息，返回是否包含，如果包含的话返回日期
    """
    for item in attraction_info:
        if any(sight in sight_item for sight_item in item['景点']):
            return item
    return None


def copy_row(source_row, target_row):
    """将源行的内容和格式复制到目标行"""
    for src_cell, tgt_cell in zip(source_row.cells, target_row.cells):
        # 复制文本
        tgt_cell.text = src_cell.text

        # 复制段落格式（如字体、大小、加粗等）
        source_para = src_cell.paragraphs[0]
        target_para = tgt_cell.paragraphs[0]

        # 复制段落样式
        target_para.style = source_para.style

        # 复制字体属性
        source_font = source_para.runs[0].font
        target_font = target_para.runs[0].font
        target_font.name = source_font.name
        target_font.size = source_font.size
        target_font.bold = source_font.bold
        target_font.italic = source_font.italic


def copy_row_and_insert(table, source_row_index, copy_count):
    """
    复制源行并插入到表格中
    """
    source_row = table.rows[source_row_index]

    for i in range(copy_count):  # 插入新行

        new_row = table.add_row()
        # 获取新行和源行的XML元素
        new_row_element = new_row._element
        source_row_element = source_row._element
        # 将新行插入到源行的位置（源行会自动后移）
        source_row_element.getparent().insert(source_row_index+1, new_row_element)

        # 复制源行的内容和格式到新行
        copy_row(source_row, table.rows[source_row_index])


def clear_table(target_table, start_row, end_row):
    """
    清空表格中指定行的数据
    """
    for i in range(start_row, end_row):
        for cell in target_table.rows[i].cells:
            cell.text = ""


def replace_in_docx(input_path, output_path, replacements):
    """
    在Word文档中执行批量替换
    """
    doc = Document(input_path)

    def replace_text(container):
        # 处理跨Run的文本替换
        for paragraph in container.paragraphs:
            merged_text = ''.join([run.text for run in paragraph.runs])
            if any(old in merged_text for old in replacements):
                # 执行批量替换
                for old, new in replacements.items():
                    merged_text = merged_text.replace(old, new)

                # 清除原有Run
                for run in paragraph.runs:
                    run.text = ''

                # 添加新Run并保留第一个Run的格式
                if paragraph.runs:
                    paragraph.runs[0].text = merged_text
                else:
                    paragraph.add_run(merged_text)

        # 处理表格中的嵌套表格
        if hasattr(container, "tables"):
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text(cell)

    # 处理文档各组成部分
    replace_text(doc)

    # 处理页眉页脚
    for section in doc.sections:
        replace_text(section.header)
        replace_text(section.footer)

    doc.save(output_path)
